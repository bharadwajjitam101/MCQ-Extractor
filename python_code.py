import pytesseract
from PIL import Image
import pandas as pd
from groq import Groq
import os
import tempfile
from llama_index.core import SimpleDirectoryReader
import re
import json
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from docx import Document
from docx.shared import Inches

# Set your Groq API key
groq_api_key = 'Add your generated groq API'

def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text

def extract_text_from_pdf(file_path):
    reader = SimpleDirectoryReader(input_files=[file_path])
    documents = reader.load_data()
    if documents:
        return " ".join([doc.text for doc in documents])
    else:
        return ""

def parse_and_format_mcqs_with_groq(text):
    client = Groq(api_key=groq_api_key)
    response = client.chat.completions.create(
        model="llama-3.1-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant that extracts and formats multiple-choice questions (MCQs) from text."
            },
            {
                "role": "user",
                "content": f"""Extract the questions and options from the following text, and format them as a list of questions with their options. Each question should be on a new line, followed by its options (A, B, C, D) on separate lines. Here's the text:

{text}

Format the output as follows:
Question: [Question text]
A) [Option A]
B) [Option B]
C) [Option C]
D) [Option D]

Respond only with the formatted questions and options, no additional text."""
            }
        ],
        temperature=0.25,
        max_tokens=1024,
        top_p=1,
        stream=False,
        stop=None
    )
    
    return response.choices[0].message.content

def process_raw_mcq_text(raw_text):
    pattern = r'Question: (.*?)\nA\) (.*?)\nB\) (.*?)\nC\) (.*?)\nD\) (.*?)(?:\n|$)'
    matches = re.findall(pattern, raw_text, re.DOTALL)
    
    mcq_data = []
    for match in matches:
        question = match[0].strip()
        options = [opt.strip() for opt in match[1:]]
        mcq_data.append({"question": question, "options": options})
    
    return mcq_data

def extract_mcqs(file_path):
    _, file_extension = os.path.splitext(file_path)
    if file_extension.lower() in ['.png', '.jpg', '.jpeg']:
        text = extract_text_from_image(file_path)
    elif file_extension.lower() == '.pdf':
        text = extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

    if text:
        print(f"Extracted text length: {len(text)} characters")
        
        max_chunk_length = 4000
        chunks = [text[i:i+max_chunk_length] for i in range(0, len(text), max_chunk_length)]
        
        print(f"Number of chunks to process: {len(chunks)}")
        
        all_mcq_data = []
        for i, chunk in enumerate(chunks):
            print(f"Processing chunk {i+1}/{len(chunks)}")
            raw_mcq_text = parse_and_format_mcqs_with_groq(chunk)
            print("Raw MCQ text from Groq:")
            print(raw_mcq_text)
            mcq_data = process_raw_mcq_text(raw_mcq_text)
            all_mcq_data.extend(mcq_data)
        
        if all_mcq_data:
            df = pd.DataFrame([
                [q['question']] + q['options']
                for q in all_mcq_data
            ], columns=["Question", "Option A", "Option B", "Option C", "Option D"])
            
            return df
        else:
            print("No MCQs found in the document.")
            return None
    else:
        print("No text found in the document.")
        return None

def save_as_csv(df, output_path):
    df.to_csv(output_path, index=False)
    print(f"CSV file saved to {output_path}")

def save_as_json(df, output_path):
    df.to_json(output_path, orient='records')
    print(f"JSON file saved to {output_path}")

def save_as_pdf(df, output_path):
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    data = [df.columns.tolist()] + df.values.tolist()
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), '#CCCCCC'),
        ('TEXTCOLOR', (0, 0), (-1, 0), '#000000'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), '#EEEEEE'),
        ('TEXTCOLOR', (0, 1), (-1, -1), '#000000'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, '#000000'),
    ]))
    pdf.build([table])
    print(f"PDF file saved to {output_path}")

def save_as_docx(df, output_path):
    doc = Document()
    doc.add_heading('MCQ Questions', 0)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = column_name

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.save(output_path)
    print(f"Word document saved to {output_path}")

def edit_mcqs(df):
    while True:
        print("\nCurrent MCQs:")
        for index, row in df.iterrows():
            print(f"{index + 1}. {row['Question']}")
        
        choice = input("\nEnter the number of the MCQ you want to edit (or 'q' to finish editing): ")
        if choice.lower() == 'q':
            break
        
        try:
            index = int(choice) - 1
            if 0 <= index < len(df):
                print("\nCurrent question and answers:")
                print(f"Question: {df.loc[index, 'Question']}")
                for option in ['A', 'B', 'C', 'D']:
                    print(f"Option {option}: {df.loc[index, f'Option {option}']}")
                
                edit_choice = input("\nWhat would you like to edit? (q/a/b/c/d): ").lower()
                if edit_choice == 'q':
                    df.loc[index, 'Question'] = input("Enter the new question: ")
                elif edit_choice in ['a', 'b', 'c', 'd']:
                    df.loc[index, f'Option {edit_choice.upper()}'] = input(f"Enter the new option {edit_choice.upper()}: ")
                else:
                    print("Invalid choice. No changes made.")
            else:
                print("Invalid MCQ number.")
        except ValueError:
            print("Invalid input. Please enter a number or 'q'.")
    
    return df

def save_files(df):
    output_formats = {
        'csv': save_as_csv,
        'json': save_as_json,
        'pdf': save_as_pdf,
        'docx': save_as_docx
    }
    
    while True:
        print("\nAvailable file formats:")
        for format in output_formats.keys():
            print(f"- {format}")
        
        choice = input("Enter the file format you want to save (or 'q' to finish): ").lower()
        if choice == 'q':
            break
        
        if choice in output_formats:
            output_path = f"mcqs.{choice}"
            output_formats[choice](df, output_path)
        else:
            print("Invalid format. Please choose from the available options.")

def main():
    file_path = input("Enter the path to your image or PDF file: ")
    df = extract_mcqs(file_path)
    
    if df is not None:
        print("\nExtracted MCQs:")
        print(df)
        
        edit_choice = input("\nWould you like to edit the MCQs? (y/n): ").lower()
        if edit_choice == 'y':
            df = edit_mcqs(df)
        
        save_files(df)
        print("Processing complete.")
    else:
        print("No MCQs were extracted. Please check your input file and try again.")

if __name__ == "__main__":
    main()